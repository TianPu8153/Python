from docx import Document
import os

OLDPATH = "C:\\Users\\Admin\\Desktop\\python\\old"
PATH = "C:\\Users\\Admin\\Desktop\\python\\new"

DICT = {
  "郭锋": "郭鑫",
  "安昊": "丁泽",
  "刘晓霖": "邢伟岭",
  "李扬": "鲍鹏举",
  "付晓蓉": "马竞楠",
  "刘宇梅": "郭烨红",
  "韩士准": "李莹",
  "刘达志": "赵宇",
  "赵飞": "王宇宁",
  "原伟": "陈晓宇",
  "刘改琴": "冯鑫",
  "杜立志":"刘益甫",
  "陈曼":"樊文婷",
  "侯丽蓉":"刘嘉俊",
  "程文龙":"苏鹏", 
  "贡国玉":"曹娜",
  "程利":"刘胜利",
  "2016":"2018",
  "2015":"2017",
  "TyCloud云管理平台":"MIMS通信机房维护巡检管理系统项目",
  "三期":"一期",
  "河北天翼科贸发展有限公司":"内蒙古电力集团蒙电信息通信产业有限责任公司"
}

def main():
    for fileName in os.listdir(OLDPATH):
        oldFile = OLDPATH + "\\" + fileName
        newFile = PATH + "\\" + fileName
        if oldFile.split(".")[1] == 'docx':
            document = Document(oldFile)
            document = check(document)
            document.save(newFile)   

def check(document):
    # tables
    for table in document.tables:
        for row in range(len(table.rows)):
            for col in range(len(table.columns)):
                for key, value in DICT.items():
                    if key in table.cell(row ,col).text:
                        print(key+"->"+value)
                        table.cell(row ,col).text = table.cell(row ,col).text.replace(key, value)

    # paragraphs
    for para in document.paragraphs:
        for i in range(len(para.runs)):
            for key, value in DICT.items():
                if key in para.runs[i].text:
                    print(key+"->"+value)
                    para.runs[i].text = para.runs[i].text.replace(key, value)

    return document 
    






if __name__ == '__main__':
	main()